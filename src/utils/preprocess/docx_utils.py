"""
Utility functions for extracting text from Microsoft Word documents.
"""
import docx
from werkzeug.datastructures import FileStorage
import io

def extract_text_from_docx(file: FileStorage) -> str:
    """
    Extract text from a Microsoft Word (.docx) document.
    
    Args:
        file (FileStorage): The uploaded Word document file
        
    Returns:
        str: Extracted text as a single string (lowercased)
    """
    try:
        file.seek(0)
        doc = docx.Document(file)
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text).lower()
    except Exception as e:
        # Optionally log the error here
        return ""
